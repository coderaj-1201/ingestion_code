"""
DOCX Parser
===========

Uses python-docx natively — no OCR needed.
Word documents have rich structure: Heading 1/2/3, Table, paragraph styles.
Header/footer are separate XML sections — removed cleanly.

Strategy:
  - Walk document.paragraphs in order
  - Use paragraph.style.name to detect Heading 1/2/3/4/Title
  - Tables extracted atomically → LLM NL summary
  - Parent-child: flush parent on each Heading 1/2 boundary
"""
from __future__ import annotations

import logging
from datetime import datetime, timezone
from uuid import uuid4

import docx
from docx.oxml.ns import qn

from processors.pdf_parser import (
    _llm_clean_page,
    _llm_serialise_table
)
from shared.config import settings
from shared.models import ChunkType, RawChunk

logger = logging.getLogger(__name__)

_HEADING_STYLES = {
    "Title":     "title",
    "Heading 1": "h1",
    "Heading 2": "h2",
    "Heading 3": "h3",
    "Heading 4": "h4",
}


def _table_to_markdown_docx(table: docx.table.Table) -> str:
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def _estimate_page(para_index: int, total_paragraphs: int, total_pages_estimate: int = 1) -> int:
    """
    Word doesn't expose page numbers in python-docx without rendering.
    Estimate based on paragraph position. Good enough for metadata.
    """
    if total_paragraphs == 0:
        return 1
    return max(1, round((para_index / total_paragraphs) * total_pages_estimate) + 1)


def parse_docx(
    file_bytes: bytes,
    doc_name: str,
    doc_url: str,
    domain: str,
    blob_path: str,
) -> list[RawChunk]:
    import io
    ingested_at = datetime.now(timezone.utc).isoformat()
    document    = docx.Document(io.BytesIO(file_bytes))

    # Estimate page count from section properties (approximate)
    total_pages = max(1, len(document.paragraphs) // 30)

    doc_title          = doc_name.replace(".docx", "").replace("_", " ")
    current_heading    = ""
    current_subheading = ""
    current_parent_id  = str(uuid4())
    current_parent_content: list[str] = []
    current_parent_page = 1
    chunks: list[RawChunk] = []

    def _flush_parent():
        nonlocal current_parent_id, current_parent_content
        if not current_parent_content:
            return
        parent = RawChunk(
            chunk_id           = current_parent_id,
            parent_id          = "",
            chunk_type         = ChunkType.HEADING if current_heading else ChunkType.PARAGRAPH,
            domain             = domain,
            doc_name           = doc_name,
            source             = doc_name,
            doc_url            = doc_url,
            file_type          = "docx",
            blob_path          = blob_path,
            ingested_at        = ingested_at,
            page_number        = current_parent_page,
            title              = doc_title,
            section_heading    = current_heading,
            section_subheading = current_subheading,
            content            = "\n\n".join(current_parent_content),
        )
        chunks.append(parent)
        current_parent_id      = str(uuid4())
        current_parent_content = []

    # Collect all block-level elements in order (paragraphs + tables)
    body = document.element.body
    para_count = len(document.paragraphs)
    para_idx   = 0

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # ── Paragraph ─────────────────────────────────────────────────────────
        if tag == "p":
            para = docx.text.paragraph.Paragraph(child, document)
            style_name = para.style.name if para.style else "Normal"
            text       = para.text.strip()
            page_num   = _estimate_page(para_idx, para_count, total_pages)
            para_idx  += 1

            if not text:
                continue

            heading_level = _HEADING_STYLES.get(style_name)

            if heading_level == "title":
                doc_title = text
                _flush_parent()
                current_parent_page = page_num
                current_heading     = text
                current_subheading  = ""

            elif heading_level in ("h1", "h2"):
                _flush_parent()
                current_parent_page = page_num
                current_heading     = text
                current_subheading  = ""
                chunks.append(RawChunk(
                    chunk_id           = str(uuid4()),
                    parent_id          = current_parent_id,
                    chunk_type         = ChunkType.HEADING,
                    domain             = domain,
                    doc_name           = doc_name,
                    source             = doc_name,
                    doc_url            = doc_url,
                    file_type          = "docx",
                    blob_path          = blob_path,
                    ingested_at        = ingested_at,
                    page_number        = page_num,
                    title              = doc_title,
                    section_heading    = current_heading,
                    section_subheading = current_subheading,
                    content            = text,
                ))
                current_parent_content.append(text)

            elif heading_level in ("h3", "h4"):
                current_subheading = text
                current_parent_content.append(text)

            else:
                # Regular paragraph — light LLM clean
                cleaned = _llm_clean_page(text, page_num)
                if not cleaned:
                    continue
                current_parent_content.append(cleaned)
                chunks.append(RawChunk(
                    chunk_id           = str(uuid4()),
                    parent_id          = current_parent_id,
                    chunk_type         = ChunkType.PARAGRAPH,
                    domain             = domain,
                    doc_name           = doc_name,
                    source             = doc_name,
                    doc_url            = doc_url,
                    file_type          = "docx",
                    blob_path          = blob_path,
                    ingested_at        = ingested_at,
                    page_number        = page_num,
                    title              = doc_title,
                    section_heading    = current_heading,
                    section_subheading = current_subheading,
                    content            = cleaned,
                ))

        # ── Table ─────────────────────────────────────────────────────────────
        elif tag == "tbl":
            table    = docx.table.Table(child, document)
            page_num = _estimate_page(para_idx, para_count, total_pages)
            tbl_md   = _table_to_markdown_docx(table)
            if not tbl_md:
                continue
            nl_summary = _llm_serialise_table(tbl_md, current_heading)
            chunks.append(RawChunk(
                chunk_id           = str(uuid4()),
                parent_id          = current_parent_id,
                chunk_type         = ChunkType.TABLE,
                domain             = domain,
                doc_name           = doc_name,
                source             = doc_name,
                doc_url            = doc_url,
                file_type          = "docx",
                blob_path          = blob_path,
                ingested_at        = ingested_at,
                page_number        = page_num,
                title              = doc_title,
                section_heading    = current_heading,
                section_subheading = current_subheading,
                content            = nl_summary,
                table_raw          = tbl_md,
            ))
            current_parent_content.append(nl_summary)

    _flush_parent()
    logger.info("DOCX parsed: %s → %d chunks", doc_name, len(chunks))
    return chunks
